"""
Post-process the ICoICT paper .docx to insert real, editable Word equations (OMML)
into Section III (Methodology).

Rationale:
    Mathematical formulations of the methods and evaluation metrics belong in the
    Methodology section of an IEEE paper, not in Results. This script converts LaTeX
    to native Office Math (OMML) via pandoc and inserts numbered equations together
    with the supporting narrative that defines each symbol.

Usage:
    python3 add_equations.py
"""
import os
import copy
import tempfile
from typing import List

import pypandoc
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

DOCX = "/projects/sandbox/MetodPen-ICoICT-/ICoICT_Paper_NGBoost_Water_Quality.docx"

FONT = "Times New Roman"
BODY_PT = Pt(10)

# Two-tab layout: a centre tab to centre the equation across the usable text width
# (A4 21cm - 1.9cm - 1.9cm = 17.2cm ~= 6.77in) and a right tab for the equation number.
CENTER_TAB = Inches(3.35)
RIGHT_TAB = Inches(6.7)


def latex_to_omath(latex: str):
    """Convert a LaTeX expression to a deep-copied OMML <m:oMath> element."""
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        pypandoc.convert_text(
            f"$$ {latex} $$", "docx", format="markdown", outputfile=tmp,
            extra_args=["--quiet"],
        )
        d = Document(tmp)
        p = d.paragraphs[0]._p
        omath_para = p.find(qn("m:oMathPara"))
        omath = omath_para.find(qn("m:oMath")) if omath_para is not None else p.find(qn("m:oMath"))
        if omath is None:
            raise RuntimeError(f"No OMML produced for: {latex}")
        return copy.deepcopy(omath)
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


def make_body(doc: Document, text: str):
    """Create a justified body paragraph matching the document style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = BODY_PT
    run.font.name = FONT
    p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    return p


def make_equation(doc: Document, latex: str, number: int):
    """Create a centred equation paragraph with a right-aligned equation number."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.tab_stops.add_tab_stop(CENTER_TAB, WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    # leading tab -> jumps to the centre tab stop
    p.add_run("\t")
    # native Office Math object
    p._p.append(latex_to_omath(latex))
    # trailing tab -> jumps to the right tab stop, followed by the equation number
    r = p.add_run(f"\t({number})")
    r.font.name = FONT
    r.font.size = BODY_PT
    return p


def relocate_after(anchor_p, new_paragraphs: List):
    """Move newly created paragraphs to sit immediately after the anchor, in order."""
    cursor = anchor_p
    for para in new_paragraphs:
        cursor.addnext(para._p)
        cursor = para._p


def find_anchor(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise RuntimeError(f"Anchor not found: {needle!r}")


def main():
    doc = Document(DOCX)
    n = 0  # equation counter

    # ---------------- III-C: Model Architectures ----------------
    anchor_c = find_anchor(doc, "Three gradient boosting variants are evaluated")
    block_c = []

    block_c.append(make_body(doc,
        "Gradient boosting constructs an additive model by sequentially fitting base "
        "learners to the functional gradient of a differentiable loss. After M boosting "
        "iterations, the prediction is expressed as:"))
    n += 1
    block_c.append(make_equation(doc, r"F_M(x) = F_0(x) + \sum_{m=1}^{M} \eta\, h_m(x)", n))
    block_c.append(make_body(doc,
        "where F_0(x) is the initial estimate, h_m(x) is the m-th base learner (a decision "
        "tree), and \u03b7 is the learning rate. Unlike conventional gradient boosting that "
        "predicts a point estimate, NGBoost predicts the parameters \u03b8 of a full conditional "
        "distribution P_\u03b8(y|x). For binary potability classification, a Bernoulli "
        "distribution is assumed with probability mass function:"))
    n += 1
    block_c.append(make_equation(doc, r"P(y \mid p) = p^{\,y}\,(1-p)^{\,1-y}, \quad y \in \{0,1\}", n))
    block_c.append(make_body(doc,
        "where p is the predicted probability of potability. NGBoost is trained by minimizing "
        "a proper scoring rule; using the logarithmic score, the loss for a single observation "
        "is the negative log-likelihood:"))
    n += 1
    block_c.append(make_equation(doc, r"S(\theta, y) = -\log P_{\theta}(y)", n))
    block_c.append(make_body(doc,
        "Proper scoring rules attain their expected minimum only when the predicted "
        "distribution equals the true data-generating distribution, which directly "
        "encourages well-calibrated probabilities. The defining feature of NGBoost is the use "
        "of the natural gradient in place of the ordinary gradient. The natural gradient "
        "rescales the ordinary gradient by the inverse Fisher information matrix I_S(\u03b8), "
        "accounting for the information geometry of the parameter space:"))
    n += 1
    block_c.append(make_equation(doc,
        r"I_S(\theta) = \mathbb{E}_{y \sim P_\theta}\!\left[\nabla_\theta \log P_\theta(y)\,"
        r"\nabla_\theta \log P_\theta(y)^{\top}\right]", n))
    n += 1
    block_c.append(make_equation(doc,
        r"\tilde{\nabla}\,S(\theta, y) = I_S(\theta)^{-1}\,\nabla_\theta S(\theta, y)", n))
    block_c.append(make_body(doc,
        "The distributional parameters are then updated as \u03b8 \u2190 \u03b8 \u2212 \u03b7 \u00b7 "
        "natural-gradient, which yields invariance to reparameterization and more stable "
        "optimization of probabilistic outputs than first-order gradient descent."))

    relocate_after(anchor_c._p, block_c)

    # ---------------- III-D: Evaluation Metrics ----------------
    anchor_d = find_anchor(doc, "Model evaluation encompasses three categories")
    block_d = []

    block_d.append(make_body(doc,
        "Let TP, TN, FP, and FN denote the numbers of true positives, true negatives, false "
        "positives, and false negatives, respectively. The discriminative metrics are defined as:"))
    n += 1
    block_d.append(make_equation(doc, r"\mathrm{Accuracy} = \frac{TP + TN}{TP + TN + FP + FN}", n))
    n += 1
    block_d.append(make_equation(doc, r"\mathrm{Precision} = \frac{TP}{TP + FP}", n))
    n += 1
    block_d.append(make_equation(doc, r"\mathrm{Recall} = \frac{TP}{TP + FN}", n))
    n += 1
    block_d.append(make_equation(doc,
        r"F_1 = 2 \cdot \frac{\mathrm{Precision} \cdot \mathrm{Recall}}"
        r"{\mathrm{Precision} + \mathrm{Recall}}", n))
    block_d.append(make_body(doc,
        "Calibration is quantified using the Negative Log-Likelihood (NLL) over N test samples "
        "and the Expected Calibration Error (ECE) computed over M equal-width probability bins:"))
    n += 1
    block_d.append(make_equation(doc,
        r"\mathrm{NLL} = -\frac{1}{N}\sum_{i=1}^{N}\left[y_i \log p_i + (1 - y_i)\log(1 - p_i)\right]", n))
    n += 1
    block_d.append(make_equation(doc,
        r"\mathrm{ECE} = \sum_{m=1}^{M} \frac{|B_m|}{N}\,\bigl|\,\mathrm{acc}(B_m) - \mathrm{conf}(B_m)\,\bigr|", n))
    block_d.append(make_body(doc,
        "where B_m is the set of samples whose predicted confidence falls in bin m, and "
        "acc(B_m) and conf(B_m) are the average accuracy and average predicted confidence within "
        "that bin. Statistical significance of accuracy differences between two models is assessed "
        "using McNemar's test with continuity correction:"))
    n += 1
    block_d.append(make_equation(doc, r"\chi^2 = \frac{(\,|b - c| - 1\,)^2}{b + c}", n))
    block_d.append(make_body(doc,
        "where b and c are the counts of samples misclassified by exactly one of the two compared "
        "models. Under the null hypothesis of equal error rates, the statistic follows a "
        "chi-squared distribution with one degree of freedom."))

    relocate_after(anchor_d._p, block_d)

    # ---------------- III-F: SMOTE-ENN Analysis ----------------
    anchor_f = find_anchor(doc, "To investigate the impact of class imbalance correction")
    block_f = []
    block_f.append(make_body(doc,
        "Formally, a synthetic minority sample is generated by interpolating between a minority "
        "instance x_i and one of its k nearest minority-class neighbors x_nn:"))
    n += 1
    block_f.append(make_equation(doc,
        r"x_{\mathrm{new}} = x_i + \delta\,(x_{nn} - x_i), \quad \delta \sim U(0,1)", n))
    block_f.append(make_body(doc,
        "where \u03b4 is a random scalar drawn from a uniform distribution on [0,1]. The Edited "
        "Nearest Neighbors step subsequently removes any instance whose class label disagrees with "
        "the majority of its k nearest neighbors, cleaning overlapping regions of the feature space."))

    relocate_after(anchor_f._p, block_f)

    doc.save(DOCX)
    print(f"Inserted {n} equations into {DOCX}")


if __name__ == "__main__":
    main()
